"""Extract text from sampled PDF/DOCX files into JSONL.

Output schema (one record per line):
    {
      "id": "first-50-chars-of-filename",
      "path": "KHAO THI/1. ABC/QT_*.pdf",
      "dept": "KHAO THI",
      "process": "1. XAY DUNG VA QUAN LY...",
      "ext": ".pdf",
      "npages": 5,
      "text": "Toàn bộ nội dung...",
      "text_len": 12345,
      "extract_status": "ok" | "empty" | "error",
      "extract_error": null | "..."
    }

Run: python scripts/02_extract_text.py
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

import pdfplumber
from docx import Document
from tqdm import tqdm

from university_quality_os.config import settings
from university_quality_os.utils import setup_logging, write_jsonl, slugify

log = setup_logging(settings.log_level)


def extract_pdf(path: Path) -> tuple[str, int]:
    """Extract text from PDF. Returns (text, num_pages)."""
    try:
        with pdfplumber.open(path) as pdf:
            pages = [(p.extract_text() or "") for p in pdf.pages]
            return "\n\n".join(pages).strip(), len(pdf.pages)
    except Exception as e:
        raise RuntimeError(f"PDF extract failed: {e}")


def extract_docx(path: Path) -> tuple[str, int]:
    """Extract text from DOCX. Returns (text, num_paragraphs)."""
    try:
        doc = Document(path)
        # Paragraphs
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    paras.append(row_text)
        return "\n".join(paras), len(doc.paragraphs)
    except Exception as e:
        raise RuntimeError(f"DOCX extract failed: {e}")


def extract_one(path: Path, src_root: Path) -> dict:
    """Extract one file → record dict."""
    rel_path = path.relative_to(src_root)
    parts = rel_path.parts

    record = {
        "id": slugify(path.stem, max_len=80),
        "path": str(rel_path).replace("\\", "/"),
        "dept": parts[0] if len(parts) > 0 else "",
        "process": parts[1] if len(parts) > 1 else "",
        "filename": path.name,
        "ext": path.suffix.lower(),
        "npages": 0,
        "text": "",
        "text_len": 0,
        "extract_status": "error",
        "extract_error": None,
    }

    try:
        if record["ext"] == ".pdf":
            text, npages = extract_pdf(path)
        elif record["ext"] == ".docx":
            text, npages = extract_docx(path)
        else:
            record["extract_status"] = "skipped"
            record["extract_error"] = f"Unsupported extension: {record['ext']}"
            return record

        record["text"] = text
        record["text_len"] = len(text)
        record["npages"] = npages
        record["extract_status"] = "ok" if text.strip() else "empty"
        if not text.strip():
            record["extract_error"] = "Extraction returned empty text (possibly image-only PDF)"
    except Exception as e:
        record["extract_error"] = str(e)
        log.warning(f"Failed: {rel_path} → {e}")

    return record


def main():
    src = settings.raw_iso_dir
    out_path = settings.extracted_dir / "iso_texts.jsonl"

    if not src.exists() or not any(src.iterdir()):
        log.error(f"No raw ISO files in {src}")
        log.error("Run scripts/01_sample_docs.py first.")
        sys.exit(1)

    # Collect all files
    files = []
    for ext in ("*.pdf", "*.PDF", "*.docx", "*.DOCX"):
        files.extend(src.rglob(ext))
    files = sorted(set(files))

    log.info(f"Found {len(files)} files in {src}")
    log.info(f"Extracting → {out_path}")
    log.info("")

    records = []
    for f in tqdm(files, desc="Extracting", unit="file"):
        rec = extract_one(f, src)
        records.append(rec)

    write_jsonl(records, out_path)

    # Stats
    ok = sum(1 for r in records if r["extract_status"] == "ok")
    empty = sum(1 for r in records if r["extract_status"] == "empty")
    error = sum(1 for r in records if r["extract_status"] == "error")
    skipped = sum(1 for r in records if r["extract_status"] == "skipped")
    avg_len = (sum(r["text_len"] for r in records if r["extract_status"] == "ok")
               / max(ok, 1))

    log.info("")
    log.info(f"📊 Extraction summary:")
    log.info(f"  ✅ OK:      {ok:4d}")
    log.info(f"  ⚠️  Empty:   {empty:4d} (likely scanned PDFs — needs OCR)")
    log.info(f"  ❌ Error:   {error:4d}")
    log.info(f"  ⏭️  Skipped: {skipped:4d}")
    log.info(f"  📏 Avg text length (OK only): {avg_len:.0f} chars")
    log.info("")
    log.info(f"📁 Output: {out_path}")

    if ok < len(files) * 0.7:
        log.warning("Less than 70% of files extracted successfully.")
        log.warning("If many empty: PDFs may be image-scanned. OCR is a follow-up task.")


if __name__ == "__main__":
    main()
