"""Sinh Word .docx template cho import CTĐT.

Mirror cấu trúc Excel template (gen_import_template.py) — 11 tables theo đúng
thứ tự, mỗi table có header row + 1 example row.

Usage:
    python scripts/gen_word_template.py --out ../../import_templates/CTDT_template.docx
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTIONS = [
    {
        "name": "00_Program",
        "title": "1. Thông tin chương trình đào tạo",
        "headers": ["code", "name_vn", "name_en", "level", "duration_years",
                    "total_credits", "language", "decision_no", "decision_date",
                    "issuing_authority"],
        "example": ["7480201", "Công nghệ thông tin", "Information Technology",
                    "DAI_HOC", "4", "144", "Tiếng Việt", "346/QĐ-ĐHKTĐN",
                    "2024-06-25", "Trường ĐH Kiến trúc Đà Nẵng"],
    },
    {
        "name": "01_PO",
        "title": "2. Mục tiêu đào tạo (PO)",
        "headers": ["code", "text_vn", "text_en", "order"],
        "example": ["PO1", "Đào tạo nguồn nhân lực có phẩm chất...", "", "1"],
    },
    {
        "name": "02_PLO",
        "title": "3. Chuẩn đầu ra (PLO)",
        "headers": ["code", "text_vn", "text_en", "order"],
        "example": ["PLO1", "Áp dụng các kiến thức cơ bản...", "", "1"],
    },
    {
        "name": "03_PI",
        "title": "4. Performance Indicator (PI)",
        "headers": ["code", "plo_code", "text_vn", "text_en", "order"],
        "example": ["PI1.1", "PLO1", "Vận dụng các kiến thức...", "", "1"],
    },
    {
        "name": "04_PLO_PO_matrix",
        "title": "5. Ma trận PLO × PO (cell = X hoặc trống)",
        "headers": ["plo_code", "PO1", "PO2", "PO3", "PO4", "PO5", "PO6"],
        "example": ["PLO1", "X", "", "", "", "", ""],
    },
    {
        "name": "05_PLO_VQF_matrix",
        "title": "6. Ma trận PLO × VQF (Khung trình độ QG VN)",
        "headers": ["plo_code",
                    "K1", "K2", "K3", "K4", "K5",
                    "S1", "S2", "S3", "S4", "S5", "S6",
                    "A1", "A2", "A3", "A4"],
        "example": ["PLO1", "X", "", "", "", "", "", "", "", "", "", "", "", "", "", ""],
    },
    {
        "name": "06_Course",
        "title": "7. Học phần",
        "headers": ["code", "name_vn", "name_en", "credits",
                    "hours_lt", "hours_th", "hours_self",
                    "prerequisites", "corequisites",
                    "knowledge_group", "is_elective", "semester_default",
                    "language", "description"],
        "example": ["MAT101", "Toán cao cấp 1", "Calculus 1", "3",
                    "30", "15", "90", "", "",
                    "DAI_CUONG", "False", "1",
                    "Tiếng Việt", "Học phần cung cấp..."],
    },
    {
        "name": "07_CLO",
        "title": "8. CLO (Chuẩn đầu ra học phần)",
        "headers": ["course_code", "clo_code", "text_vn", "text_en", "order"],
        "example": ["MAT101", "CLO1", "Tính được giới hạn...", "", "1"],
    },
    {
        "name": "08_CLO_PI_matrix",
        "title": "9. Ma trận CLO × PI (level: I/R/M/A)",
        "headers": ["course_code", "clo_code", "pi_code", "level"],
        "example": ["MAT101", "CLO1", "PI1.1", "I"],
    },
    {
        "name": "09_Assessment",
        "title": "10. Cấu phần đánh giá",
        "headers": ["course_code", "component_name", "weight_pct",
                    "method", "clo_codes", "order"],
        "example": ["MAT101", "Thi cuối kỳ", "60", "Tự luận 90p",
                    "CLO1,CLO2,CLO3", "3"],
    },
    {
        "name": "10_WeeklyPlan",
        "title": "11. Kế hoạch giảng dạy theo tuần",
        "headers": ["course_code", "week", "topic", "content",
                    "hours_lt", "hours_th", "clo_codes"],
        "example": ["MAT101", "1", "Giới hạn dãy số",
                    "Định nghĩa, các tính chất...", "3", "0", "CLO1"],
    },
]


def make_doc() -> Document:
    doc = Document()
    title = doc.add_heading("CĐR Steward — Template Import CTĐT (Word)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "Template chứa 11 bảng theo thứ tự chuẩn. Hệ thống parse theo THỨ TỰ "
        "(không theo tên section), nên KHÔNG đảo trật tự hay xóa bảng. "
        "Có thể xóa dòng ví dụ, thay bằng dữ liệu thực."
    ).italic = True
    doc.add_paragraph()

    rules = doc.add_paragraph()
    rules.add_run("Lưu ý:\n").bold = True
    rules.add_run(
        "• Header (dòng đầu mỗi bảng) — KHÔNG sửa tên cột.\n"
        "• Ma trận: cell điền 'X' hoặc để trống.\n"
        "• CLO_PI matrix: level chỉ chấp nhận I / R / M / A.\n"
        "• Tổng weight_pct mỗi học phần ở bảng Assessment phải = 100.\n"
        "• Date format: YYYY-MM-DD (vd 2024-06-25).\n"
        "• Enum hợp lệ: level = DAI_HOC|THAC_SI|TIEN_SI; "
        "knowledge_group = DAI_CUONG|CO_SO|CHUYEN_NGANH|TU_CHON|TOT_NGHIEP.\n"
    )
    doc.add_paragraph()

    for section in SECTIONS:
        h = doc.add_heading(f"{section['title']}  [{section['name']}]", 1)

        table = doc.add_table(rows=2, cols=len(section["headers"]))
        table.style = "Light Grid Accent 1"

        for i, header in enumerate(section["headers"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for i, val in enumerate(section["example"]):
            table.rows[1].cells[i].text = val

        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph(
        "Sau khi điền xong, upload qua tab 'Import Excel/Word' trên web app "
        "hoặc CLI: python scripts/import_docx_cli.py --file <path>.docx",
        style="Intense Quote",
    )

    return doc


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--out", required=True)
    args = p.parse_args()

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = make_doc()
    doc.save(out_path)
    print(f"[OK] wrote {out_path.resolve()}")
    print(f"     11 tables ({sum(len(s['headers']) for s in SECTIONS)} columns total)")


if __name__ == "__main__":
    main()
